"""File service to encapsulate files' manipulations."""

import io
import logging
import os
import re
import subprocess
import tempfile
from contextlib import contextmanager
from pathlib import Path

import mutagen
from docx import Document
from docx.shared import Pt, RGBColor
from minio import Minio
from minio.error import MinioException, S3Error

from summary.core.config import get_settings

settings = get_settings()


logger = logging.getLogger(__name__)


class FileServiceException(Exception):
    """Base exception for file service operations."""

    pass


class FileService:
    """Service for downloading and preparing files from MinIO storage."""

    def __init__(self):
        """Initialize FileService with MinIO client and configuration."""
        endpoint = (
            settings.aws_s3_endpoint_url.removeprefix("https://")
            .removeprefix("http://")
            .rstrip("/")
        )

        self._minio_client = Minio(
            endpoint,
            access_key=settings.aws_s3_access_key_id,
            secret_key=settings.aws_s3_secret_access_key.get_secret_value(),
            secure=settings.aws_s3_secure_access,
        )

        self._bucket_name = settings.aws_storage_bucket_name
        self._stream_chunk_size = 32 * 1024

        self._allowed_extensions = settings.recording_allowed_extensions
        self._max_duration = settings.recording_max_duration

    def markdown_to_docx(self, content: str, title: str = "") -> bytes:
        """Convert markdown transcription content to a docx document.

        Handles: H1/H2 headings (# / ##), timestamp lines ([MM:SS]), plain paragraphs.

        Returns raw bytes of the .docx file.
        """
        doc = Document()

        if title:
            heading = doc.add_heading(title, level=1)
            heading.runs[0].font.size = Pt(18)
            heading.runs[0].font.color.rgb = RGBColor(0x4A, 0x3C, 0x5C)

        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif re.match(r"^\[\d{2}:\d{2}\]", stripped):
                # Timestamp line — monospace-style paragraph
                para = doc.add_paragraph(stripped)
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            else:
                doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def upload_to_minio(self, object_key: str, content: str | bytes, content_type: str = "text/markdown") -> None:
        """Upload text content to MinIO storage.

        Args:
            object_key: Destination path in MinIO (e.g. "transcriptions/{uuid}.md").
            content: Text content to upload.
            content_type: MIME type of the content.

        Raises:
            FileServiceException: If the upload fails.
        """
        logger.info("Uploading to MinIO | object_key: %s", object_key)

        data = content if isinstance(content, bytes) else content.encode("utf-8")
        try:
            self._minio_client.put_object(
                self._bucket_name,
                object_key,
                io.BytesIO(data),
                length=len(data),
                content_type=content_type,
            )
            logger.info("Upload successful | object_key: %s", object_key)
        except (MinioException, S3Error) as e:
            raise FileServiceException(
                f"Failed to upload {object_key}"
            ) from e

    def _download_from_minio(self, remote_object_key) -> Path:
        """Download file from MinIO to local temporary file.

        The file is downloaded to a temporary location for local manipulation
        such as validation, conversion, or processing before being used.
        """
        logger.info("Download recording | object_key: %s", remote_object_key)

        if not remote_object_key:
            logger.warning("Invalid object_key '%s'", remote_object_key)
            raise ValueError("Invalid object_key")

        extension = Path(remote_object_key).suffix.lower()

        if extension not in self._allowed_extensions:
            logger.warning("Invalid file extension '%s'", extension)
            raise ValueError(f"Invalid file extension '{extension}'")

        response = None

        try:
            response = self._minio_client.get_object(
                self._bucket_name, remote_object_key
            )

            with tempfile.NamedTemporaryFile(
                suffix=extension, delete=False, prefix="minio_download_"
            ) as tmp:
                for chunk in response.stream(self._stream_chunk_size):
                    tmp.write(chunk)

                tmp.flush()
                local_path = Path(tmp.name)

                logger.info("Recording successfully downloaded")
                logger.debug("Recording local file path: %s", local_path)

                return local_path

        except (MinioException, S3Error) as e:
            raise FileServiceException(
                "Unexpected error while downloading object."
            ) from e

        finally:
            if response:
                response.close()

    def _validate_duration(self, local_path: Path) -> float:
        """Validate audio file duration against configured maximum."""
        file_metadata = mutagen.File(local_path).info
        duration = file_metadata.length

        logger.info(
            "Recording file duration: %.2f seconds",
            duration,
        )

        if self._max_duration is not None and duration > self._max_duration:
            error_msg = "Recording too long. Limit is %.2fs seconds" % (
                self._max_duration,
            )
            logger.error(error_msg)
            raise ValueError(error_msg)

        return duration

    def _extract_audio_from_video(self, video_path: Path) -> Path:
        """Extract audio from video file (e.g., MP4) and save as audio file."""
        logger.info("Extracting audio from video file: %s", video_path)

        with tempfile.NamedTemporaryFile(
            suffix=".m4a", delete=False, prefix="audio_extract_"
        ) as tmp:
            output_path = Path(tmp.name)

        try:
            command = [
                "ffmpeg",
                "-i",
                str(video_path),
                "-vn",  # No video
                "-acodec",
                "copy",
                "-y",  # Overwrite output file if exists
                str(output_path),
            ]

            # ruff: noqa: S603
            subprocess.run(
                command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
            )

            logger.info("Audio successfully extracted to: %s", output_path)
            return output_path

        except FileNotFoundError as e:
            logger.error("ffmpeg not found. Please install ffmpeg.")
            if output_path.exists():
                os.remove(output_path)
            raise RuntimeError("ffmpeg is not installed or not in PATH") from e
        except subprocess.CalledProcessError as e:
            logger.error("Audio extraction failed: %s", e.stderr.decode())
            if output_path.exists():
                os.remove(output_path)
            raise RuntimeError("Failed to extract audio.") from e

    def _preprocess_audio(self, input_path: Path) -> Path:
        """Normalize audio for whisper: highpass filter + EBU loudness normalization + 16kHz mono.

        highpass=f=80 removes sub-80Hz room noise (AC hum, rumble) without touching voice.
        loudnorm=I=-16 applies EBU R128 normalization to compensate level variance between speakers.
        16kHz mono is whisper's native input format, avoids internal resampling inside the model.
        """
        logger.info("Preprocessing audio: highpass + loudnorm + 16kHz mono")

        with tempfile.NamedTemporaryFile(
            suffix=".wav", delete=False, prefix="audio_preprocessed_"
        ) as tmp:
            output_path = Path(tmp.name)

        try:
            command = [
                "ffmpeg",
                "-i",
                str(input_path),
                "-af",
                "highpass=f=80,loudnorm=I=-16:LRA=11:TP=-1.5",
                "-ar",
                "16000",
                "-ac",
                "1",
                "-y",
                str(output_path),
            ]

            # ruff: noqa: S603
            subprocess.run(
                command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
            )

            logger.info("Audio preprocessed successfully: %s", output_path)
            return output_path

        except FileNotFoundError as e:
            logger.error("ffmpeg not found. Please install ffmpeg.")
            if output_path.exists():
                os.remove(output_path)
            raise RuntimeError("ffmpeg is not installed or not in PATH") from e
        except subprocess.CalledProcessError as e:
            logger.error("Audio preprocessing failed: %s", e.stderr.decode())
            if output_path.exists():
                os.remove(output_path)
            raise RuntimeError("Failed to preprocess audio.") from e

    @contextmanager
    def prepare_audio_file(self, remote_object_key: str):
        """Download and prepare audio file for processing.

        Downloads file from MinIO, validates duration, and yields an open
        file handle with metadata. Automatically cleans up temporary files
        when the context exits.
        """
        downloaded_path = None
        processed_path = None
        preprocessed_path = None
        file_handle = None

        try:
            downloaded_path = self._download_from_minio(remote_object_key)
            duration = self._validate_duration(downloaded_path)

            extension = downloaded_path.suffix.lower()

            if extension in settings.recording_video_extensions:
                logger.info("Video file detected, extracting audio...")
                extracted_audio_path = self._extract_audio_from_video(downloaded_path)
                processed_path = extracted_audio_path
            else:
                processed_path = downloaded_path

            preprocessed_path = self._preprocess_audio(processed_path)

            metadata = {"duration": duration, "extension": extension}

            file_handle = open(preprocessed_path, "rb")
            yield file_handle, metadata

        finally:
            if file_handle:
                file_handle.close()

            for path in [downloaded_path, processed_path, preprocessed_path]:
                if path is None or not os.path.exists(path):
                    continue

                try:
                    os.remove(path)
                    logger.debug("Temporary file removed: %s", path)
                except OSError as e:
                    logger.warning("Failed to remove temporary file %s: %s", path, e)
